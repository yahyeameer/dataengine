"""
The client-facing renderings of a report: PDF, Word and Excel.

`report.py` decides what a month-end report says. This decides what it looks
like when it leaves the building with a firm's name on it.

Three formats rather than one, because the reader does three different things
with it. The PDF is what gets emailed to the client and read on a phone -- it
is the one that has to look like a document somebody was paid to produce. The
Word file is what a partner edits before sending, adding the paragraph about
the thing only they know. The workbook is what the client's own finance person
opens to check a figure against their system, so its numbers are real numbers
in real cells, not a picture of numbers.

**No new system dependencies.** reportlab and python-docx are pure-Python
wheels and openpyxl was already here, so the agent image gains three pip
packages and no apt packages. That matters more than it sounds: the VPS this
runs on has one core, and the obvious alternative -- rendering HTML to PDF
through a browser or a Cairo/Pango stack -- is a hundred megabytes of system
libraries and a rendering process that competes with the parser for the only
core there is.

**Branding is one colour and one name.** `Brand.for_client` takes an accent and
derives the rest -- the deep tone for headers, the tint behind a card, and
whether the text on the accent should be white or near-black. A firm that gives
us `#8A1538` gets a document that looks designed for it rather than a blue
document with a maroon rectangle. Anything unreadable or malformed falls back
to the product's own palette rather than failing the job: a report that arrives
in the wrong blue is a complaint, and a report that does not arrive is an
incident.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Any

from .report import (
    Bars,
    Callout,
    Footnote,
    Heading,
    KeyFigures,
    Prose,
    ReportDocument,
    Table,
    render_markdown,
)

#: Everything `render_document` can produce. `md` is here so that one caller
#: can hold one list and not special-case the default.
FORMATS = ("md", "pdf", "docx", "xlsx")

CONTENT_TYPES = {
    "md": "text/markdown; charset=utf-8",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


# ---------------------------------------------------------------------------
# Colour
# ---------------------------------------------------------------------------

_HEX = re.compile(r"^#?(?:[0-9a-fA-F]{3}|[0-9a-fA-F]{6})$")


def _hex(value: Any, fallback: str) -> str:
    """A validated `#rrggbb`, or the fallback. Never raises: see the docstring."""
    if not isinstance(value, str) or not _HEX.match(value.strip()):
        return fallback
    digits = value.strip().lstrip("#").lower()
    if len(digits) == 3:
        digits = "".join(character * 2 for character in digits)
    return f"#{digits}"


def _rgb(value: str) -> tuple[int, int, int]:
    digits = value.lstrip("#")
    return (int(digits[0:2], 16), int(digits[2:4], 16), int(digits[4:6], 16))


def _mix(first: str, second: str, weight: float) -> str:
    """`weight` is how much of `second` to take, 0..1."""
    a, b = _rgb(first), _rgb(second)
    blended = tuple(round(a[i] + (b[i] - a[i]) * weight) for i in range(3))
    return "#{:02x}{:02x}{:02x}".format(*blended)


def _relative_luminance(value: str) -> float:
    """WCAG relative luminance, used only to decide white-or-dark ink."""
    channels = []
    for raw in _rgb(value):
        c = raw / 255
        channels.append(c / 12.92 if c <= 0.03928 else ((c + 0.055) / 1.055) ** 2.4)
    return 0.2126 * channels[0] + 0.7152 * channels[1] + 0.0722 * channels[2]


def _ink_on(background: str) -> str:
    """
    The text colour that survives on this background.

    A brand colour is chosen to look good behind a logo, not to carry 9pt type,
    and pale ones are common -- a bank's light gold, a lettings agency's mint.
    White on those is unreadable, so the band flips to near-black instead of
    quietly failing the contrast the rest of the document is careful about.
    """
    return "#ffffff" if _relative_luminance(background) < 0.45 else "#10131a"


@dataclass(frozen=True)
class Brand:
    """
    Everything the three renderers need to know about whose document this is.

    Built through `for_client`, which is the only place the derivations live.
    """

    name: str = "DataEngine"
    accent: str = "#1f5fbf"
    accent_deep: str = "#143f80"
    accent_ink: str = "#ffffff"
    tint: str = "#f1f5fc"
    ink: str = "#10131a"
    muted: str = "#5b6473"
    rule: str = "#d9dee6"
    negative: str = "#b3261e"
    positive: str = "#12805a"
    warning: str = "#9a6700"
    footer: str = ""

    @classmethod
    def for_client(
        cls,
        *,
        name: str | None = None,
        accent: Any = None,
        footer: str | None = None,
    ) -> "Brand":
        base = _hex(accent, cls.accent)
        return cls(
            # A firm's name can be anything a person typed into a form. Trimmed
            # and truncated because it is drawn into a fixed band, and a
            # 300-character organisation name would run off the page rather
            # than wrap.
            name=(name or cls.name).strip()[:60] or cls.name,
            accent=base,
            accent_deep=_mix(base, "#000000", 0.32),
            accent_ink=_ink_on(base),
            tint=_mix(base, "#ffffff", 0.93),
            footer=(footer or "").strip()[:120],
        )

    def tone(self, name: str) -> str:
        return {"danger": self.negative, "warning": self.warning, "info": self.accent}.get(
            name, self.accent
        )


# ---------------------------------------------------------------------------
# Values
# ---------------------------------------------------------------------------

_NEGATIVE = re.compile(r"^\((.*)\)$")
_NUMERIC = re.compile(r"^[-+]?[\d,]+(?:\.\d+)?$")


def _is_negative(text: str) -> bool:
    """`report._money` writes credits in parentheses, as a set of accounts does."""
    return bool(_NEGATIVE.match(text.strip()))


def _as_number(text: str) -> float | None:
    """
    The number behind a formatted cell, or None if there isn't one.

    Only Excel asks for this, and it is the difference between a workbook the
    client's finance person can sum and a workbook they have to retype. The
    parentheses convention has to be undone here: `(£150.00)` is -150.0, and
    writing it as text would leave a column that looks like money and cannot be
    added up.
    """
    stripped = text.strip()
    if not stripped or stripped == "—":
        return None

    negative = bool(_NEGATIVE.match(stripped))
    if negative:
        stripped = stripped[1:-1]

    stripped = re.sub(r"[£$€\s]", "", stripped)
    if not _NUMERIC.match(stripped):
        return None

    try:
        value = float(stripped.replace(",", ""))
    except ValueError:
        return None
    return -value if negative else value


def _currency_symbol(text: str) -> str | None:
    for symbol in ("£", "$", "€"):
        if symbol in text:
            return symbol
    return None


def _is_definition_table(block: Table) -> bool:
    """
    A table with no column names is a list of facts, not a table.

    `report.py` builds the provenance trail as `Table(["", ""], ...)` -- source
    workbook, uploaded, derived from -- and drawing that with a header row
    gives it a solid bar of brand colour containing nothing at all. It also
    wants its second column wide, because the values there are filenames and
    timestamps rather than amounts.
    """
    return len(block.headers) == 2 and not any(header.strip() for header in block.headers)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def render_pdf(document: ReportDocument, brand: Brand) -> bytes:
    """
    The copy that gets emailed to the client.

    Helvetica throughout, deliberately. The product's own typeface is IBM Plex,
    and embedding it would mean shipping four font files in the agent image and
    keeping their licence straight for documents that leave our control. The
    base-14 fonts are in every PDF reader ever written, so the file is small and
    renders identically on a phone, in a browser preview and in whatever their
    accountant opens attachments with.
    """
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        BaseDocTemplate,
        Flowable,
        Frame,
        KeepTogether,
        NextPageTemplate,
        PageTemplate,
        Paragraph,
        Spacer,
    )
    from reportlab.platypus import Table as PdfTable
    from reportlab.platypus import TableStyle
    from reportlab.platypus.flowables import HRFlowable

    page_width, page_height = A4
    margin = 18 * mm
    band_height = 24 * mm
    content_width = page_width - 2 * margin

    accent = colors.HexColor(brand.accent)
    accent_deep = colors.HexColor(brand.accent_deep)
    accent_ink = colors.HexColor(brand.accent_ink)
    ink = colors.HexColor(brand.ink)
    muted = colors.HexColor(brand.muted)
    rule = colors.HexColor(brand.rule)
    tint = colors.HexColor(brand.tint)
    negative = colors.HexColor(brand.negative)

    body = ParagraphStyle(
        "body", fontName="Helvetica", fontSize=9.5, leading=14, textColor=ink, spaceAfter=6
    )
    small = ParagraphStyle("small", parent=body, fontSize=8.2, leading=11.5, textColor=muted)
    title = ParagraphStyle(
        "title", fontName="Helvetica-Bold", fontSize=21, leading=25, textColor=ink, spaceAfter=2
    )
    subtitle = ParagraphStyle("subtitle", parent=body, fontSize=9, textColor=muted, spaceAfter=2)
    heading = ParagraphStyle(
        "heading",
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        textColor=accent_deep,
        spaceBefore=15,
        spaceAfter=5,
        keepWithNext=1,
    )
    cell = ParagraphStyle("cell", parent=body, fontSize=8.6, leading=11.5, spaceAfter=0)
    cell_right = ParagraphStyle("cell_right", parent=cell, alignment=TA_RIGHT)
    cell_head = ParagraphStyle(
        "cell_head", parent=cell, fontName="Helvetica-Bold", textColor=accent_ink
    )
    cell_head_right = ParagraphStyle("cell_head_right", parent=cell_head, alignment=TA_RIGHT)
    card_label = ParagraphStyle(
        "card_label", parent=body, fontSize=7.4, leading=9, textColor=muted, spaceAfter=3
    )
    card_value = ParagraphStyle(
        "card_value", fontName="Helvetica-Bold", fontSize=14.5, leading=17, textColor=ink
    )
    card_note = ParagraphStyle("card_note", parent=card_label, fontSize=7, spaceAfter=0)

    class _Bars(Flowable):
        """
        A ranked bar chart, drawn straight onto the canvas.

        Not a reportlab chart object: those are built for axes and legends, and
        this is a list with a length attached to each row. Drawing it directly
        is fewer lines than configuring one, and it lets the bar carry the sign
        of the value -- a credit is red and reads as an exception at a glance,
        which is the entire reason anyone scans this section.
        """

        def __init__(self, items: list[tuple[str, float]], labels: list[str]) -> None:
            super().__init__()
            self.items = items
            self.labels = labels
            self.row_height = 15
            self.label_width = 42 * mm
            self.value_width = 30 * mm

        def wrap(self, available_width: float, _available_height: float):
            self.width = available_width
            self.height = len(self.items) * self.row_height
            return self.width, self.height

        def draw(self) -> None:
            canvas = self.canv
            largest = max((abs(value) for _label, value in self.items), default=0) or 1.0
            track = max(self.width - self.label_width - self.value_width, 10)
            y = self.height - self.row_height

            for (label, value), text in zip(self.items, self.labels):
                canvas.setFont("Helvetica", 8.4)
                canvas.setFillColor(muted)
                canvas.drawString(0, y + 4.5, label[:30])

                canvas.setFillColor(tint)
                canvas.roundRect(self.label_width, y + 2, track, self.row_height - 6, 2, 0, 1)

                filled = max(abs(value) / largest * track, 1.2)
                canvas.setFillColor(negative if value < 0 else accent)
                canvas.roundRect(self.label_width, y + 2, filled, self.row_height - 6, 2, 0, 1)

                canvas.setFont("Helvetica-Bold", 8.4)
                canvas.setFillColor(negative if value < 0 else ink)
                canvas.drawRightString(self.width, y + 4.5, text)
                y -= self.row_height

    def _cell(text: str, style: ParagraphStyle, colour_negative: bool = False) -> Paragraph:
        if colour_negative and _is_negative(text):
            return Paragraph(f'<font color="{brand.negative}">{_escape(text)}</font>', style)
        return Paragraph(_escape(text), style)

    def _table(block: Table) -> PdfTable:
        numeric = set(block.numeric)
        columns = len(block.headers)
        definition = _is_definition_table(block)

        if definition:
            # Two columns of prose: the value side holds a filename or a
            # timestamp, which wraps to two lines in a money-sized column.
            label_width = content_width * 0.40
            figure_width = content_width - label_width
        else:
            # The label column earns the space; figures need only enough to hold
            # the widest amount, and a table of three equal columns wastes the
            # page on whitespace between a total and its heading.
            figure_width = min(30 * mm, content_width / max(columns, 1))
            label_width = content_width - figure_width * (columns - 1)

        data = []
        if not definition:
            data.append(
                [
                    _cell(header, cell_head_right if index in numeric else cell_head)
                    for index, header in enumerate(block.headers)
                ]
            )
        data.extend(
            [
                _cell(
                    value,
                    cell_right if index in numeric else cell,
                    colour_negative=index in numeric,
                )
                for index, value in enumerate(row)
            ]
            for row in block.rows
        )

        first_body_row = 0 if definition else 1
        table = PdfTable(
            data,
            colWidths=[label_width] + [figure_width] * (columns - 1),
            repeatRows=0 if definition else 1,
            hAlign="LEFT",
        )
        commands = [
            ("ROWBACKGROUNDS", (0, first_body_row), (-1, -1), [colors.white, tint]),
            ("LINEBELOW", (0, first_body_row), (-1, -1), 0.4, rule),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 7),
            ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ]
        if not definition:
            commands.insert(0, ("BACKGROUND", (0, 0), (-1, 0), accent_deep))
        table.setStyle(TableStyle(commands))
        return table

    def _cards(block: KeyFigures) -> list:
        """Three to a row: four is too narrow for a formatted total at 14pt."""
        rows: list = []
        figures = list(block.figures)
        per_row = 3
        gap = 5 * mm
        card_width = (content_width - gap * (per_row - 1)) / per_row

        for start in range(0, len(figures), per_row):
            group = figures[start : start + per_row]
            cells = []
            for figure in group:
                stack = [Paragraph(_escape(figure.label).upper(), card_label)]
                # A total is short and wants to be big; a date range is long
                # and wraps to three lines at the same size, which drags every
                # card in the row down with it.
                style = (
                    card_value
                    if len(figure.value) <= 14
                    else ParagraphStyle("card_value_long", parent=card_value, fontSize=11, leading=14)
                )
                stack.append(_cell(figure.value, style, colour_negative=True))
                if figure.note:
                    stack.append(
                        Paragraph(_escape(f"{figure.note[0]} {figure.note[1]}"), card_note)
                    )
                cells.append(stack)

            # Spacer columns rather than cell padding: a gap that is part of the
            # table would be tinted like the card beside it.
            interleaved: list = []
            widths: list = []
            for index, content in enumerate(cells):
                if index:
                    interleaved.append("")
                    widths.append(gap)
                interleaved.append(content)
                widths.append(card_width)

            card = PdfTable([interleaved], colWidths=widths, hAlign="LEFT")
            style = [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
            ]
            for index in range(len(cells)):
                column = index * 2
                style.append(("BACKGROUND", (column, 0), (column, 0), tint))
                style.append(("LINEBEFORE", (column, 0), (column, 0), 2.2, accent))
            for index in range(1, len(cells)):
                column = index * 2 - 1
                style.append(("LEFTPADDING", (column, 0), (column, 0), 0))
                style.append(("RIGHTPADDING", (column, 0), (column, 0), 0))
            card.setStyle(TableStyle(style))
            rows.append(card)
            rows.append(Spacer(1, gap))

        return rows

    def _callout(block: Callout) -> PdfTable:
        colour = colors.HexColor(brand.tone(block.tone))
        stack = [
            Paragraph(
                f'<b><font color="{brand.tone(block.tone)}">{_escape(block.title)}</font></b>',
                body,
            )
        ]
        stack.extend(Paragraph(_escape(line), small) for line in block.lines)
        table = PdfTable([[stack]], colWidths=[content_width], hAlign="LEFT")
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(_mix(brand.tone(block.tone), "#ffffff", 0.92))),
                    ("LINEBEFORE", (0, 0), (0, -1), 2.5, colour),
                    ("TOPPADDING", (0, 0), (-1, -1), 9),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]
            )
        )
        return table

    # --- page furniture ----------------------------------------------------

    def _band(canvas, _doc) -> None:
        canvas.saveState()
        canvas.setFillColor(accent)
        canvas.rect(0, page_height - band_height, page_width, band_height, stroke=0, fill=1)
        canvas.setFillColor(accent_ink)
        canvas.setFont("Helvetica-Bold", 13)
        canvas.drawString(margin, page_height - band_height + 9.5 * mm, brand.name)
        canvas.setFont("Helvetica", 8.5)
        canvas.drawRightString(
            page_width - margin, page_height - band_height + 9.9 * mm, "Month-end report"
        )
        canvas.restoreState()
        _footer(canvas, _doc)

    def _running_head(canvas, _doc) -> None:
        canvas.saveState()
        canvas.setFillColor(muted)
        canvas.setFont("Helvetica", 8)
        canvas.drawString(margin, page_height - 14 * mm, f"{brand.name} · {document.title}")
        canvas.setStrokeColor(rule)
        canvas.setLineWidth(0.5)
        canvas.line(margin, page_height - 16 * mm, page_width - margin, page_height - 16 * mm)
        canvas.restoreState()
        _footer(canvas, _doc)

    def _footer(canvas, _doc) -> None:
        canvas.saveState()
        canvas.setStrokeColor(rule)
        canvas.setLineWidth(0.5)
        canvas.line(margin, 15 * mm, page_width - margin, 15 * mm)
        canvas.setFillColor(muted)
        canvas.setFont("Helvetica", 7.5)
        canvas.drawString(margin, 11 * mm, brand.footer or brand.name)
        canvas.restoreState()

    class _Numbered(_numbered_canvas_base()):
        """Page x of y, which needs the total and therefore a second pass."""

        def _stamp(self, total: int) -> None:
            self.saveState()
            self.setFillColor(muted)
            self.setFont("Helvetica", 7.5)
            self.drawRightString(page_width - margin, 11 * mm, f"Page {self._pageNumber} of {total}")
            self.restoreState()

    buffer = io.BytesIO()
    pdf = BaseDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=band_height + 8 * mm,
        bottomMargin=20 * mm,
        title=f"{document.title} — {document.workspace_name}",
        author=brand.name,
        subject=f"Dataset version {document.version_no}",
    )
    first = Frame(
        margin,
        20 * mm,
        content_width,
        page_height - band_height - 8 * mm - 20 * mm,
        id="first",
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    later = Frame(
        margin,
        20 * mm,
        content_width,
        page_height - 22 * mm - 20 * mm,
        id="later",
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    pdf.addPageTemplates(
        [
            PageTemplate(id="first", frames=[first], onPage=_band),
            PageTemplate(id="later", frames=[later], onPage=_running_head),
        ]
    )

    story: list = [
        NextPageTemplate("later"),
        Paragraph(_escape(document.title), title),
        Paragraph(
            _escape(
                f"{document.workspace_name} · dataset version {document.version_no} · "
                f"generated {document.generated_at.strftime('%d %B %Y %H:%M UTC')}"
            ),
            subtitle,
        ),
        HRFlowable(width="100%", thickness=1.6, color=accent, spaceBefore=8, spaceAfter=4),
    ]

    for block in document.blocks:
        if isinstance(block, Heading):
            story.append(Paragraph(_escape(block.text), heading))
            story.append(
                HRFlowable(width="100%", thickness=0.5, color=rule, spaceAfter=8, spaceBefore=0)
            )
        elif isinstance(block, Prose):
            story.append(Paragraph(_escape(block.text), body))
        elif isinstance(block, Callout):
            story.append(_callout(block))
            story.append(Spacer(1, 6 * mm))
        elif isinstance(block, KeyFigures):
            story.extend(_cards(block))
        elif isinstance(block, Table):
            story.append(_table(block))
            story.append(Spacer(1, 4 * mm))
        elif isinstance(block, Bars):
            story.append(KeepTogether(_Bars(block.items, block.labels)))
            story.append(Spacer(1, 4 * mm))
        elif isinstance(block, Footnote):
            story.append(Spacer(1, 4 * mm))
            story.append(HRFlowable(width="100%", thickness=0.5, color=rule, spaceAfter=6))
            story.append(Paragraph(_escape(block.text), small))

    pdf.build(story, canvasmaker=_Numbered)
    return buffer.getvalue()


def _numbered_canvas_base():
    """
    The two-pass canvas, as a factory so the import stays inside the PDF path.

    Every page is held back until the whole document is built, at which point
    the page count is known and each held page is stamped and emitted. This is
    the standard reportlab recipe; it is here rather than inline because a
    class statement at module scope would import reportlab on any import of
    this module, including the ones that only want Excel.
    """
    from reportlab.pdfgen import canvas as pdfcanvas

    class NumberedCanvas(pdfcanvas.Canvas):
        def __init__(self, *args: Any, **kwargs: Any) -> None:
            super().__init__(*args, **kwargs)
            self._pages: list[dict[str, Any]] = []

        def showPage(self) -> None:  # noqa: N802 - reportlab's name
            self._pages.append(dict(self.__dict__))
            self._startPage()

        def save(self) -> None:
            total = len(self._pages)
            for state in self._pages:
                self.__dict__.update(state)
                self._stamp(total)
                super().showPage()
            super().save()

        def _stamp(self, total: int) -> None:  # overridden by the subclass above
            pass

    return NumberedCanvas


#: Characters the base-14 fonts cannot draw, and what to draw instead.
#:
#: Helvetica is WinAnsi-encoded, so an arrow renders as a black box rather than
#: as an arrow -- and `→` is not exotic here: `report.py` writes every period
#: comparison as "£163,900.50 → £184,320.55". The narrative is worse, because a
#: language model will happily reach for an en-dash-arrow or a mathematical
#: symbol in a sentence nobody proofreads.
_PDF_SUBSTITUTIONS = {
    "→": "->",
    "←": "<-",
    "↑": "^",
    "↓": "v",
    "≥": ">=",
    "≤": "<=",
    "≈": "~",
    "█": "|",
    "░": ".",
}


def _escape(text: str) -> str:
    """
    reportlab's Paragraph reads a small HTML dialect, so client data is markup
    until it is escaped. A vendor called `Smith & Sons <Holdings>` is not
    hypothetical, and the failure is a paragraph that raises mid-build and
    takes the whole report with it.

    The encode/decode round trip afterwards is the font's limit, not ours:
    anything WinAnsi cannot represent becomes `?` here rather than a black box
    on a document with a client's name on it.
    """
    value = str(text)
    for character, replacement in _PDF_SUBSTITUTIONS.items():
        value = value.replace(character, replacement)
    value = value.encode("cp1252", "replace").decode("cp1252")
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------


def render_docx(document: ReportDocument, brand: Brand) -> bytes:
    """
    The copy a partner edits before it goes out.

    Everything is a real Word construct -- styles, tables, shading, a page-number
    field -- rather than a picture of one, because the entire reason this format
    exists is that somebody is going to change a sentence in it.
    """
    from docx import Document as DocxDocument
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Emu, Inches, Pt, RGBColor, Twips

    def colour(value: str) -> RGBColor:
        return RGBColor(*_rgb(value))

    def fill(cell: Any, value: str) -> None:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), value.lstrip("#"))
        cell._tc.get_or_add_tcPr().append(shading)

    def borders(cell: Any, **edges: tuple[str, int]) -> None:
        """edges: left=("1f5fbf", 24) -- colour and eighths of a point."""
        properties = cell._tc.get_or_add_tcPr()
        element = OxmlElement("w:tcBorders")
        for edge, (value, size) in edges.items():
            side = OxmlElement(f"w:{edge}")
            side.set(qn("w:val"), "single")
            side.set(qn("w:sz"), str(size))
            side.set(qn("w:color"), value.lstrip("#"))
            element.append(side)
        properties.append(element)

    def fixed(table: Any) -> None:
        table.autofit = False
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        table._tbl.tblPr.append(layout)

    def margins(table: Any, top: int, right: int, bottom: int, left: int) -> None:
        """
        Cell padding, in twentieths of a point.

        Word puts none above or below a cell's text by default, so a table with
        the paragraph spacing turned off -- which is the only way to stop Word's
        8pt-after from doubling every row -- ends up with rows whose text touches
        the rule above it. Padding belongs on the table, where it can differ
        between a data row and a figure card.
        """
        element = OxmlElement("w:tblCellMar")
        for edge, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")
            element.append(node)
        table._tbl.tblPr.append(element)

    def no_borders(table: Any) -> None:
        element = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            side = OxmlElement(f"w:{edge}")
            side.set(qn("w:val"), "none")
            side.set(qn("w:sz"), "0")
            element.append(side)
        table._tbl.tblPr.append(element)

    def run(paragraph: Any, text: str, *, size=9.5, bold=False, ink=None, caps=False):
        item = paragraph.add_run(text.upper() if caps else text)
        item.font.size = Pt(size)
        item.bold = bold
        item.font.color.rgb = colour(ink or brand.ink)
        return item

    doc = DocxDocument()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = colour(brand.ink)
    # Word's default template puts 8pt after every paragraph, including the one
    # inside every table cell. Left on, each table row carries an invisible
    # blank line and the document reads as a draft. Spacing is set deliberately
    # below, wherever it is wanted.
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.08
    # Word resolves East Asian and complex-script faces separately; without
    # this the document opens in Calibri on one machine and in whatever the
    # theme says on another.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    usable = section.page_width - section.left_margin - section.right_margin

    # --- the band ----------------------------------------------------------
    #
    # One cell with a right-hand tab stop, not two cells side by side. Two
    # shaded cells leave a hairline of page colour at the seam between them --
    # invisible in Word at 100%, and a white scratch down the middle of the
    # brand colour in the PDF the client is actually sent.
    band = doc.add_table(rows=1, cols=1)
    fixed(band)
    no_borders(band)
    margins(band, 150, 140, 150, 140)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band_cell = band.rows[0].cells[0]
    band_cell.width = int(usable)
    fill(band_cell, brand.accent)
    banner = band_cell.paragraphs[0]
    # Short of the full width by the cell's own margin, or the right-hand text
    # tabs past the edge and wraps onto a second line inside the band.
    banner.paragraph_format.tab_stops.add_tab_stop(
        Emu(int(usable) - int(Inches(0.3))), WD_TAB_ALIGNMENT.RIGHT
    )
    run(banner, brand.name, size=13, bold=True, ink=brand.accent_ink)
    banner.add_run("\t")
    run(banner, "Month-end report", size=8.5, ink=brand.accent_ink)

    heading_paragraph = doc.add_paragraph()
    heading_paragraph.paragraph_format.space_before = Pt(14)
    heading_paragraph.paragraph_format.space_after = Pt(0)
    run(heading_paragraph, document.title, size=20, bold=True)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    run(
        meta,
        f"{document.workspace_name} · dataset version {document.version_no} · "
        f"generated {document.generated_at.strftime('%d %B %Y %H:%M UTC')}",
        size=8.5,
        ink=brand.muted,
    )

    def rule_after(paragraph: Any, value: str, size: int = 12) -> None:
        properties = paragraph._p.get_or_add_pPr()
        element = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), value.lstrip("#"))
        element.append(bottom)
        properties.append(element)

    rule_after(meta, brand.accent, 12)

    # --- blocks ------------------------------------------------------------
    for block in document.blocks:
        if isinstance(block, Heading):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(14)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.keep_with_next = True
            run(paragraph, block.text, size=12.5, bold=True, ink=brand.accent_deep)
            rule_after(paragraph, brand.rule, 4)

        elif isinstance(block, Prose):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            run(paragraph, block.text)

        elif isinstance(block, Callout):
            tone = brand.tone(block.tone)
            table = doc.add_table(rows=1, cols=1)
            fixed(table)
            no_borders(table)
            margins(table, 120, 140, 120, 140)
            cell = table.rows[0].cells[0]
            cell.width = int(usable)
            fill(cell, _mix(tone, "#ffffff", 0.92))
            borders(cell, left=(tone, 24))
            first = cell.paragraphs[0]
            run(first, block.title, bold=True, ink=tone)
            for line in block.lines:
                paragraph = cell.add_paragraph()
                run(paragraph, line, size=8.5, ink=brand.muted)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        elif isinstance(block, KeyFigures):
            figures = list(block.figures)
            per_row = 3
            for start in range(0, len(figures), per_row):
                group = figures[start : start + per_row]
                table = doc.add_table(rows=1, cols=per_row)
                fixed(table)
                no_borders(table)
                margins(table, 110, 130, 130, 130)
                for index in range(per_row):
                    cell = table.rows[0].cells[index]
                    cell.width = int(usable / per_row)
                    if index >= len(group):
                        continue
                    figure = group[index]
                    fill(cell, brand.tint)
                    borders(cell, left=(brand.accent, 18))
                    label = cell.paragraphs[0]
                    label.paragraph_format.space_after = Pt(1)
                    run(label, figure.label, size=7.5, ink=brand.muted, caps=True)
                    value = cell.add_paragraph()
                    value.paragraph_format.space_after = Pt(0)
                    run(
                        value,
                        figure.value,
                        size=14,
                        bold=True,
                        ink=brand.negative if _is_negative(figure.value) else brand.ink,
                    )
                    if figure.note:
                        note = cell.add_paragraph()
                        note.paragraph_format.space_after = Pt(0)
                        run(note, f"{figure.note[0]} {figure.note[1]}", size=7, ink=brand.muted)
                doc.add_paragraph().paragraph_format.space_after = Pt(2)

        elif isinstance(block, Table):
            numeric = set(block.numeric)
            columns = len(block.headers)
            definition = _is_definition_table(block)
            table = doc.add_table(rows=0 if definition else 1, cols=columns)
            fixed(table)
            no_borders(table)
            margins(table, 70, 110, 70, 110)
            if definition:
                widths = [int(usable * 0.40), int(usable * 0.60)]
            else:
                figure_width = min(int(Inches(1.05)), int(usable / max(columns, 1)))
                widths = [usable - figure_width * (columns - 1)] + [figure_width] * (columns - 1)

            if not definition:
                for index, header in enumerate(block.headers):
                    cell = table.rows[0].cells[index]
                    cell.width = int(widths[index])
                    fill(cell, brand.accent_deep)
                    paragraph = cell.paragraphs[0]
                    if index in numeric:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run(paragraph, header, size=8.5, bold=True, ink=brand.accent_ink)

            for position, row in enumerate(block.rows):
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cell = cells[index]
                    cell.width = int(widths[index])
                    if position % 2:
                        fill(cell, brand.tint)
                    borders(cell, bottom=(brand.rule, 4))
                    paragraph = cell.paragraphs[0]
                    if index in numeric:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run(
                        paragraph,
                        value,
                        size=8.6,
                        ink=brand.negative if index in numeric and _is_negative(value) else brand.ink,
                    )
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        elif isinstance(block, Bars):
            largest = max((abs(value) for _label, value in block.items), default=0) or 1.0
            table = doc.add_table(rows=0, cols=3)
            fixed(table)
            no_borders(table)
            # No side padding: the bar starts exactly where its column does.
            margins(table, 40, 0, 40, 0)
            label_width = int(Inches(1.6))
            value_width = int(Inches(1.0))
            track = int(usable) - label_width - value_width

            for (label, value), text in zip(block.items, block.labels):
                cells = table.add_row().cells
                cells[0].width = label_width
                cells[1].width = track
                cells[2].width = value_width

                run(cells[0].paragraphs[0], label[:30], size=8.4, ink=brand.muted)

                # The bar is a two-cell table whose cell widths carry the value.
                # Word has no drawing primitive that survives an edit, and a
                # shaded cell does: resize the page and the bar resizes with it.
                inner = cells[1].add_table(rows=1, cols=2)
                fixed(inner)
                no_borders(inner)
                margins(inner, 0, 0, 0, 0)
                filled = max(int(track * abs(value) / largest), int(Twips(20)))
                inner.rows[0].cells[0].width = filled
                inner.rows[0].cells[1].width = max(track - filled, int(Twips(20)))
                fill(inner.rows[0].cells[0], brand.negative if value < 0 else brand.accent)
                fill(inner.rows[0].cells[1], brand.tint)
                for cell in inner.rows[0].cells:
                    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                    cell.paragraphs[0].add_run("").font.size = Pt(6)

                # A cell must *end* with a paragraph -- Word repairs one that
                # does not -- but it must not begin with the empty paragraph
                # python-docx leaves in place when a table is added to a cell.
                # That empty line is a full row of its own, and it renders every
                # bar one line below the label and the figure it belongs to.
                tail = cells[1].add_paragraph()
                tail.paragraph_format.space_after = Pt(0)
                # The size of the *paragraph mark*, not of a run. An empty
                # paragraph is as tall as its mark, so one left at the body size
                # adds a 9.5pt line under every bar -- which is most of the
                # height of the bar itself.
                size = OxmlElement("w:sz")
                size.set(qn("w:val"), "2")
                mark = OxmlElement("w:rPr")
                mark.append(size)
                # Appended last: in the pPr schema the paragraph mark's own run
                # properties come after the spacing set above.
                tail._p.get_or_add_pPr().append(mark)
                leading = cells[1].paragraphs[0]._p
                leading.getparent().remove(leading)

                for cell in (cells[0], cells[2]):
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                paragraph = cells[2].paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run(
                    paragraph,
                    text,
                    size=8.4,
                    bold=True,
                    ink=brand.negative if value < 0 else brand.ink,
                )
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        elif isinstance(block, Footnote):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(10)
            rule_after(paragraph, brand.rule, 4)
            note = doc.add_paragraph()
            italic = run(note, block.text, size=7.8, ink=brand.muted)
            italic.italic = True

    # --- footer ------------------------------------------------------------
    footer = section.footer.paragraphs[0]
    run(footer, brand.footer or brand.name, size=7.5, ink=brand.muted)
    footer.add_run("\t")
    page_run = footer.add_run()
    page_run.font.size = Pt(7.5)
    page_run.font.color.rgb = colour(brand.muted)
    for instruction, kind in (("begin", "w:fldChar"), ("PAGE", "w:instrText"), ("end", "w:fldChar")):
        element = OxmlElement(kind)
        if kind == "w:fldChar":
            element.set(qn("w:fldCharType"), instruction)
        else:
            element.set(qn("xml:space"), "preserve")
            element.text = " PAGE "
        page_run._r.append(element)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------


def render_xlsx(document: ReportDocument, brand: Brand) -> bytes:
    """
    The copy the client's own finance person opens.

    Every figure that can be a number is written as a number. A formatted
    string in a cell looks the same and cannot be summed, sorted or charted,
    and the first thing anyone does with a report in Excel is select a column
    and look at the status bar.
    """
    from openpyxl import Workbook
    from openpyxl.chart import BarChart, Reference
    from openpyxl.formatting.rule import DataBarRule
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    def fill(value: str) -> PatternFill:
        return PatternFill("solid", fgColor=value.lstrip("#").upper())

    def font(size=10, bold=False, ink=None, italic=False) -> Font:
        return Font(
            name="Calibri",
            size=size,
            bold=bold,
            italic=italic,
            color=(ink or brand.ink).lstrip("#").upper(),
        )

    thin = Side(style="thin", color=brand.rule.lstrip("#").upper())
    accent_side = Side(style="thick", color=brand.accent.lstrip("#").upper())

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Report"
    sheet.sheet_properties.tabColor = brand.accent.lstrip("#").upper()
    sheet.sheet_view.showGridLines = False
    sheet.page_setup.fitToWidth = 1
    sheet.sheet_properties.pageSetUpPr.fitToPage = True

    widths = {"A": 2, "B": 34, "C": 16, "D": 16, "E": 16, "F": 16, "G": 16}
    for column, width in widths.items():
        sheet.column_dimensions[column].width = width

    # --- the band ----------------------------------------------------------
    for column in range(1, 8):
        for row in (1, 2):
            sheet.cell(row=row, column=column).fill = fill(brand.accent)
    sheet.row_dimensions[1].height = 8
    sheet.row_dimensions[2].height = 24
    banner = sheet.cell(row=2, column=2, value=brand.name)
    banner.font = font(14, bold=True, ink=brand.accent_ink)
    banner.alignment = Alignment(vertical="center")
    corner = sheet.cell(row=2, column=6, value="Month-end report")
    corner.font = font(9, ink=brand.accent_ink)
    corner.alignment = Alignment(vertical="center", horizontal="right")
    sheet.merge_cells(start_row=2, start_column=6, end_row=2, end_column=7)

    title = sheet.cell(row=4, column=2, value=document.title)
    title.font = font(18, bold=True)
    sheet.row_dimensions[4].height = 26
    meta = sheet.cell(
        row=5,
        column=2,
        value=(
            f"{document.workspace_name} · dataset version {document.version_no} · "
            f"generated {document.generated_at.strftime('%d %B %Y %H:%M UTC')}"
        ),
    )
    meta.font = font(9, ink=brand.muted)
    sheet.freeze_panes = "A6"

    cursor = 7
    chart_placed = False

    for block in document.blocks:
        if isinstance(block, Heading):
            cell = sheet.cell(row=cursor, column=2, value=block.text)
            cell.font = font(12, bold=True, ink=brand.accent_deep)
            for column in range(2, 8):
                sheet.cell(row=cursor, column=column).border = Border(bottom=thin)
            cursor += 2

        elif isinstance(block, Prose):
            cell = sheet.cell(row=cursor, column=2, value=block.text)
            cell.font = font(10)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            sheet.merge_cells(start_row=cursor, start_column=2, end_row=cursor, end_column=7)
            # Roughly one line per 95 characters across the merged width. Merged
            # cells do not auto-size, and a summary paragraph clipped to one
            # line is the difference between a report and a receipt.
            sheet.row_dimensions[cursor].height = max(15, 14 * (len(block.text) // 95 + 1))
            cursor += 2

        elif isinstance(block, Callout):
            tone = brand.tone(block.tone)
            lines = [block.title, *block.lines]
            for index, line in enumerate(lines):
                cell = sheet.cell(row=cursor, column=2, value=line)
                cell.font = font(10, bold=index == 0, ink=tone if index == 0 else brand.ink)
                cell.alignment = Alignment(wrap_text=True, vertical="center")
                sheet.merge_cells(start_row=cursor, start_column=2, end_row=cursor, end_column=7)
                for column in range(2, 8):
                    target = sheet.cell(row=cursor, column=column)
                    target.fill = fill(_mix(tone, "#ffffff", 0.92))
                    if column == 2:
                        target.border = Border(left=accent_side)
                cursor += 1
            cursor += 1

        elif isinstance(block, KeyFigures):
            for index, figure in enumerate(block.figures):
                column = 2 + (index % 3) * 2
                row = cursor + (index // 3) * 3
                label = sheet.cell(row=row, column=column, value=figure.label.upper())
                label.font = font(8, ink=brand.muted)
                number = _as_number(figure.value)
                value_cell = sheet.cell(
                    row=row + 1,
                    column=column,
                    value=number if number is not None else figure.value,
                )
                if number is not None:
                    symbol = _currency_symbol(figure.value)
                    value_cell.number_format = (
                        f'{symbol}#,##0.00;[Red]({symbol}#,##0.00)' if symbol else "#,##0"
                    )
                value_cell.font = font(
                    14,
                    bold=True,
                    ink=brand.negative if _is_negative(figure.value) else brand.ink,
                )
                # Left, against the label above it. Excel right-aligns numbers
                # by default, which parks the figure at the far edge of the tile
                # and reads as a column that lost its heading.
                value_cell.alignment = Alignment(horizontal="left", vertical="center")
                note_cell = sheet.cell(
                    row=row + 2,
                    column=column,
                    value=f"{figure.note[0]} {figure.note[1]}" if figure.note else None,
                )
                note_cell.font = font(8, ink=brand.muted)
                for offset in range(3):
                    for span in range(2):
                        target = sheet.cell(row=row + offset, column=column + span)
                        target.fill = fill(brand.tint)
                        if span == 0:
                            target.border = Border(left=accent_side)
            rows_used = ((len(block.figures) - 1) // 3 + 1) * 3 if block.figures else 0
            cursor += rows_used + 1

        elif isinstance(block, Table):
            numeric = set(block.numeric)
            if not _is_definition_table(block):
                for index, header in enumerate(block.headers):
                    cell = sheet.cell(row=cursor, column=2 + index, value=header)
                    cell.fill = fill(brand.accent_deep)
                    cell.font = font(10, bold=True, ink=brand.accent_ink)
                    cell.alignment = Alignment(horizontal="right" if index in numeric else "left")
                cursor += 1

            for position, row in enumerate(block.rows):
                for index, value in enumerate(row):
                    number = _as_number(value) if index in numeric else None
                    cell = sheet.cell(
                        row=cursor, column=2 + index, value=number if number is not None else value
                    )
                    if number is not None:
                        symbol = _currency_symbol(value)
                        cell.number_format = (
                            f'{symbol}#,##0.00;[Red]({symbol}#,##0.00)' if symbol else "#,##0"
                        )
                    cell.font = font(10)
                    cell.alignment = Alignment(horizontal="right" if index in numeric else "left")
                    cell.border = Border(bottom=thin)
                    if position % 2:
                        cell.fill = fill(brand.tint)
                cursor += 1
            cursor += 1

        elif isinstance(block, Bars):
            start = cursor
            for (label, value), text in zip(block.items, block.labels):
                sheet.cell(row=cursor, column=2, value=label).font = font(10)
                cell = sheet.cell(row=cursor, column=3, value=value)
                symbol = _currency_symbol(text)
                cell.number_format = (
                    f'{symbol}#,##0.00;[Red]({symbol}#,##0.00)' if symbol else "#,##0.00"
                )
                cell.font = font(10, bold=True)
                cursor += 1

            if cursor > start:
                # Data bars rather than a chart for every block: they live in
                # the cells, so sorting the range keeps them, and they cannot
                # drift over the section below when a row is inserted.
                sheet.conditional_formatting.add(
                    f"C{start}:C{cursor - 1}",
                    DataBarRule(
                        start_type="num",
                        start_value=0,
                        end_type="max",
                        color=brand.accent.lstrip("#").upper(),
                        showValue=True,
                    ),
                )

                # One real chart, for the first series only. A workbook with a
                # chart per section is a workbook nobody can print.
                if not chart_placed:
                    chart = BarChart()
                    chart.type = "col"
                    chart.style = 2
                    chart.title = None
                    chart.legend = None
                    chart.height = 6
                    chart.width = 12
                    chart.y_axis.majorGridlines = None
                    # openpyxl leaves `delete` unset, and Excel reads an unset
                    # flag on a chart it did not create as "hidden" -- which
                    # produces a bar chart with no month names under the bars
                    # and no scale beside them.
                    chart.x_axis.delete = False
                    chart.y_axis.delete = False
                    data = Reference(sheet, min_col=3, min_row=start, max_row=cursor - 1)
                    categories = Reference(sheet, min_col=2, min_row=start, max_row=cursor - 1)
                    chart.add_data(data, titles_from_data=False)
                    chart.set_categories(categories)
                    series = chart.series[0]
                    series.graphicalProperties.solidFill = brand.accent.lstrip("#").upper()
                    series.graphicalProperties.line.noFill = True
                    sheet.add_chart(chart, f"E{start}")
                    chart_placed = True

            cursor += 1

        elif isinstance(block, Footnote):
            cursor += 1
            for column in range(2, 8):
                sheet.cell(row=cursor, column=column).border = Border(top=thin)
            cell = sheet.cell(row=cursor + 1, column=2, value=block.text)
            cell.font = font(8, ink=brand.muted, italic=True)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            sheet.merge_cells(start_row=cursor + 1, start_column=2, end_row=cursor + 1, end_column=7)
            sheet.row_dimensions[cursor + 1].height = 14 * (len(block.text) // 95 + 1)
            cursor += 3

    sheet.oddFooter.left.text = brand.footer or brand.name
    sheet.oddFooter.right.text = "Page &P of &N"

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------


def render_document(
    document: ReportDocument, fmt: str, brand: Brand | None = None
) -> tuple[bytes, str, str]:
    """
    One entry point: bytes, content type, extension.

    Unknown formats fall back to Markdown rather than raising. The format
    arrives from a job payload, which is to say from a browser, and a month-end
    report is not the thing to fail over a typo in a query parameter.
    """
    chosen = fmt if fmt in FORMATS else "md"
    brand = brand or Brand()

    if chosen == "pdf":
        return render_pdf(document, brand), CONTENT_TYPES["pdf"], "pdf"
    if chosen == "docx":
        return render_docx(document, brand), CONTENT_TYPES["docx"], "docx"
    if chosen == "xlsx":
        return render_xlsx(document, brand), CONTENT_TYPES["xlsx"], "xlsx"
    return render_markdown(document).encode("utf-8"), CONTENT_TYPES["md"], "md"


__all__ = [
    "CONTENT_TYPES",
    "FORMATS",
    "Brand",
    "render_docx",
    "render_document",
    "render_pdf",
    "render_xlsx",
]
